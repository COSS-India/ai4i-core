#!/usr/bin/env python3
"""
Convert markdown file to Word document
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Title (single #)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            i += 1
            continue
        
        # Heading level 2 (##)
        if line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
            i += 1
            continue
        
        # Heading level 3 (###)
        if line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=2)
            i += 1
            continue
        
        # Heading level 4 (####)
        if line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=3)
            i += 1
            continue
        
        # Horizontal rule (---)
        if line.startswith('---'):
            i += 1
            continue
        
        # Bold text (**text**)
        if '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)
            i += 1
            continue
        
        # Bullet list (- or *)
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Handle bold in list items
            if '**' in text:
                p = doc.add_paragraph(style='List Bullet')
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            else:
                doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            if '**' in text:
                p = doc.add_paragraph(style='List Number')
                parts = re.split(r'(\*\*.*?\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            else:
                doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Regular paragraph
        if line:
            # Handle bold in paragraphs
            if '**' in line:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        p.add_run(part)
            else:
                doc.add_paragraph(line)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == '__main__':
    import sys
    md_file = 'AI4I_CORE_MICROSERVICES_INTEGRATION_EFFORTS.md'
    docx_file = 'AI4I_CORE_MICROSERVICES_INTEGRATION_EFFORTS.docx'
    parse_markdown_to_docx(md_file, docx_file)


